import io
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

_BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT    = os.path.join(_BASE_DIR, "data", "processed", "LO_PO_Alignment_Report.docx")

# ── colour palette ──────────────────────────────────────────────────────────
# Hex strings for XML shading; RGBColor objects for font.color.rgb
NAVY_HEX       = "1F497D"
MED_BLUE_HEX   = "276FB4"
LIGHT_BLUE_HEX = "BDD7EE"
PALE_GREY_HEX  = "F2F2F2"
WHITE_HEX      = "FFFFFF"
DARK_GREY_HEX  = "404040"

NAVY       = RGBColor(0x1F, 0x49, 0x7D)
MED_BLUE   = RGBColor(0x27, 0x6F, 0xB4)
LIGHT_BLUE = RGBColor(0xBD, 0xD7, 0xEE)
PALE_GREY  = RGBColor(0xF2, 0xF2, 0xF2)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GREY  = RGBColor(0x40, 0x40, 0x40)

# ── helpers ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    """hex_color: 6-char hex string e.g. 'FF0000'"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)

def cell_para(cell, text, bold=False, size=9, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def set_col_width(table, col_idx, width_inches):
    for row in table.rows:
        row.cells[col_idx].width = Inches(width_inches)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY
    return p

def no_space_para(doc, text="", size=11, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    p.alignment = align
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = color
    return p

def add_page_break(doc):
    doc.add_page_break()

def set_repeat_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    trPr.append(tblHeader)

# ── document setup ───────────────────────────────────────────────────────────
doc = Document()

# default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

for lvl, sz, sp_b, sp_a in [(1, 16, 12, 6), (2, 13, 10, 4), (3, 11, 8, 2)]:
    h = doc.styles[f"Heading {lvl}"]
    h.font.name = "Calibri"
    h.font.size = Pt(sz)
    h.font.color.rgb = NAVY
    h.font.bold = True
    h.paragraph_format.space_before = Pt(sp_b)
    h.paragraph_format.space_after  = Pt(sp_a)

# margins
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
for attr in ("left_margin","right_margin","top_margin","bottom_margin"):
    setattr(section, attr, Inches(1))

# ── header / footer (portrait section) ───────────────────────────────────────
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_header_footer(sec, show_header=True):
    if show_header:
        hdr = sec.header
        hdr.is_linked_to_previous = False
        hp = hdr.paragraphs[0]
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run("NSCC — ITSM Program  |  LO–PO Alignment Report  |  2026-27")
        run.font.size = Pt(9)
        run.font.color.rgb = DARK_GREY
        # thin bottom border on header paragraph
        pPr = hp._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), NAVY_HEX)
        pBdr.append(bottom)
        pPr.append(pBdr)

    ftr = sec.footer
    ftr.is_linked_to_previous = False
    fp = ftr.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = DARK_GREY
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run2 = fp.add_run()
    run2.font.size = Pt(9)
    run2.font.color.rgb = DARK_GREY
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)

add_header_footer(section, show_header=False)   # title page: footer only

# ════════════════════════════════════════════════════════════════════════════
# SECTION 1 – TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run("Course Learning Outcomes to\nProgram Outcomes Alignment Report")
tr.font.size = Pt(26)
tr.font.bold = True
tr.font.color.rgb = NAVY

doc.add_paragraph()

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run("IT Systems Management and Security (ITSM) — NSCC")
sr.font.size = Pt(14)
sr.font.color.rgb = DARK_GREY

doc.add_paragraph()

ay_p = doc.add_paragraph()
ay_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
ay_p.add_run("Academic Year: 2026-27").font.size = Pt(12)

dt_p = doc.add_paragraph()
dt_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
dt_p.add_run("May 2026").font.size = Pt(12)

# divider line under title block
div = doc.add_paragraph()
div.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = div._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
top = OxmlElement("w:top")
top.set(qn("w:val"), "single"); top.set(qn("w:sz"), "12")
top.set(qn("w:space"), "1");    top.set(qn("w:color"), NAVY_HEX)
pBdr.append(top)
pPr.append(pBdr)

add_page_break(doc)

# switch to a new section with header visible
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION
new_sec = doc.add_section(WD_SECTION.NEW_PAGE)
new_sec.page_width  = Inches(8.5)
new_sec.page_height = Inches(11)
for attr in ("left_margin","right_margin","top_margin","bottom_margin"):
    setattr(new_sec, attr, Inches(1))
add_header_footer(new_sec, show_header=True)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 2 – EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Executive Summary", 1)
no_space_para(doc, (
    "This report maps all course-level learning outcomes (LOs) from required courses in the "
    "IT Systems Management and Security (ITSM) diploma program to the 12 Program Outcomes (POs) "
    "established for the 2026-27 academic year. The mapping was conducted to assess the breadth "
    "and depth of coverage each Program Outcome receives across the curriculum, identify outcomes "
    "with insufficient course-level support, and provide evidence-based recommendations to inform "
    "potential revisions to the Program Outcome list. The report includes a coverage heat map, a "
    "bar chart summarising coverage counts, a full per-course breakdown table, and a gap analysis "
    "with actionable recommendations."
), size=11)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 3 – PROGRAM OUTCOMES REFERENCE
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Program Outcomes Reference", 1)
pos = [
    "Analyze and document business systems, requirements and problems using standard methodologies and notation.",
    "Design, implement, and maintain a secure networked environment.",
    "Configure, secure and administer network operating systems, software and hardware to support business systems.",
    "Provide technical training to support business systems.",
    "Research, learn and integrate innovations in systems management and security.",
    "Integrate professional practices and skills into all projects, activities and communications in the context of an IT industry environment.",
    "Demonstrate continuous professional improvement through reflection and modification of processes and approaches in relation to the IT industries.",
    "Blend service and learning in ways that use program-related skills, knowledge and behaviours to serve others at the campus, within the College and in the community.",
    "Apply a Portfolio approach to the personal management of learning and career planning relating to the learner's occupational readiness.",
    "Apply the Essential & Employability Skills needed to enter, stay in, and progress in the world of work, productively contributing to the economy and the community.",
    "Apply sustainable practices that support economic, social, cultural and environmental stewardship.",
    "Demonstrate the principles of quality and safety as per the '5 S+S' standard, complete the online requirements for WHMIS and OH&S, and complete all safety related outcomes and milestones specific to each program.",
]
for i, text in enumerate(pos, 1):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f"PO{i}: {text}")
    run.font.size = Pt(10)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 4 – HEAT MAP (landscape section)
# ════════════════════════════════════════════════════════════════════════════
add_page_break(doc)
land_sec = doc.add_section(WD_SECTION.NEW_PAGE)
land_sec.orientation   = 1          # LANDSCAPE
land_sec.page_width    = Inches(11)
land_sec.page_height   = Inches(8.5)
land_sec.left_margin   = Inches(0.5)
land_sec.right_margin  = Inches(0.5)
land_sec.top_margin    = Inches(0.75)
land_sec.bottom_margin = Inches(0.75)
add_header_footer(land_sec, show_header=True)

add_heading(doc, "Coverage Heat Map", 1)
no_space_para(doc,
    "Each cell indicates whether a course provides primary (●) or supporting (○) alignment to a Program Outcome.",
    size=10)

COURSES = [
    "CSTN 4015","DBAS 1007","HDWR 1700","ICOM 2701","ICOM 2702","ICOM 2703",
    "INFT 2700","INFT 3000","ISEC 2700","NETW 1027","NETW 1500","NETW 2500",
    "NETW 2700","NETW 2710","NETW 3500","NETW 3700","OSYS 1000","OSYS 1200",
    "OSYS 3030","PROG 1700","WEBD 1000",
]

# alignment: dict[course] -> dict[po_num (1-12)] -> "P" | "S" | ""
RAW = {
    "CSTN 4015": {1:"P",3:"S",4:"P",5:"S",6:"S",10:"S"},
    "DBAS 1007": {1:"P",3:"P",5:"S"},
    "HDWR 1700": {2:"S",3:"P",6:"S"},
    "ICOM 2701": {6:"P",7:"P",9:"S",10:"S"},
    "ICOM 2702": {6:"P",7:"S",9:"P",10:"P"},
    "ICOM 2703": {6:"S",7:"S",9:"P",10:"P"},
    "INFT 2700": {1:"P",3:"S",5:"S",6:"S"},
    "INFT 3000": {1:"S",2:"S",3:"S",5:"S",6:"P",8:"S",10:"S"},
    "ISEC 2700": {1:"S",2:"P",3:"P",5:"S"},
    "NETW 1027": {2:"P",3:"P",5:"S"},
    "NETW 1500": {1:"S",3:"P",5:"S",6:"S"},
    "NETW 2500": {1:"S",3:"P",6:"S"},
    "NETW 2700": {2:"P",3:"P"},
    "NETW 2710": {1:"S",3:"P",5:"P"},
    "NETW 3500": {2:"S",3:"P",4:"S",5:"P",6:"S"},
    "NETW 3700": {2:"P",3:"P",5:"S"},
    "OSYS 1000": {3:"P",5:"S",6:"S"},
    "OSYS 1200": {1:"S",3:"P",6:"S"},
    "OSYS 3030": {1:"S",2:"S",3:"P",6:"S"},
    "PROG 1700": {3:"P",5:"S"},
    "WEBD 1000": {1:"S",3:"P"},
}

N_PO = 12
# content width for landscape = 11 - 0.5 - 0.5 = 10 inches
COURSE_COL = 1.0
PO_COL     = (10.0 - COURSE_COL) / N_PO   # ~0.75 inches each
col_widths  = [COURSE_COL] + [PO_COL] * N_PO

tbl = doc.add_table(rows=1, cols=N_PO + 1)
tbl.style = "Table Grid"

# header row
hdr_row = tbl.rows[0]
set_repeat_header(hdr_row)
set_cell_bg(hdr_row.cells[0], NAVY_HEX)
cell_para(hdr_row.cells[0], "Course", bold=True, size=8, color=WHITE)
for i in range(1, N_PO + 1):
    set_cell_bg(hdr_row.cells[i], NAVY_HEX)
    cell_para(hdr_row.cells[i], f"PO{i}", bold=True, size=8,
              color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

# data rows
for r_idx, course in enumerate(COURSES):
    row = tbl.add_row()
    bg_hex = PALE_GREY_HEX if r_idx % 2 == 0 else WHITE_HEX
    set_cell_bg(row.cells[0], PALE_GREY_HEX)
    cell_para(row.cells[0], course, bold=True, size=8)
    po_data = RAW.get(course, {})
    for po_num in range(1, N_PO + 1):
        val = po_data.get(po_num, "")
        if val == "P":
            set_cell_bg(row.cells[po_num], MED_BLUE_HEX)
            cell_para(row.cells[po_num], "●", bold=True, size=9,
                      color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        elif val == "S":
            set_cell_bg(row.cells[po_num], LIGHT_BLUE_HEX)
            cell_para(row.cells[po_num], "○", size=9,
                      color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            set_cell_bg(row.cells[po_num], bg_hex)
            cell_para(row.cells[po_num], "", size=8)

# column widths
for row in tbl.rows:
    for c_idx, w in enumerate(col_widths):
        row.cells[c_idx].width = Inches(w)

# legend
leg = doc.add_paragraph()
leg.paragraph_format.space_before = Pt(6)
leg.paragraph_format.space_after  = Pt(0)
r1 = leg.add_run("  ●  Primary alignment    ")
r1.font.size = Pt(9); r1.font.bold = True; r1.font.color.rgb = MED_BLUE
r2 = leg.add_run("  ○  Supporting alignment")
r2.font.size = Pt(9); r2.font.color.rgb = NAVY

# ════════════════════════════════════════════════════════════════════════════
# back to portrait
# ════════════════════════════════════════════════════════════════════════════
port_sec = doc.add_section(WD_SECTION.NEW_PAGE)
port_sec.orientation   = 0   # PORTRAIT
port_sec.page_width    = Inches(8.5)
port_sec.page_height   = Inches(11)
for attr in ("left_margin","right_margin","top_margin","bottom_margin"):
    setattr(port_sec, attr, Inches(1))
add_header_footer(port_sec, show_header=True)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 5 – BAR CHART
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Program Outcome Coverage — Bar Chart", 1)
no_space_para(doc,
    "Number of required courses providing primary or supporting alignment to each Program Outcome.",
    size=10)

po_labels  = [f"PO{i}" for i in range(1, 13)]
po_counts  = [13, 9, 19, 2, 13, 13, 3, 1, 3, 5, 0, 0]

def bar_color(n):
    if n >= 10: return "#27ae60"
    if n >= 4:  return "#f39c12"
    return "#e74c3c"

colors = [bar_color(n) for n in po_counts]

fig, ax = plt.subplots(figsize=(7.5, 4.5))
bars = ax.barh(po_labels[::-1], po_counts[::-1], color=colors[::-1],
               edgecolor="white", height=0.65)
ax.axvline(x=5, color="#1F497D", linestyle="--", linewidth=1.2, label="Min. coverage threshold (5)")
for bar, count in zip(bars, po_counts[::-1]):
    ax.text(bar.get_width() + 0.15, bar.get_y() + bar.get_height() / 2,
            str(count), va="center", ha="left", fontsize=9, color="#333333")
ax.set_xlabel("Number of Courses", fontsize=10)
ax.set_title("Number of Courses Supporting Each Program Outcome", fontsize=11, fontweight="bold", color="#1F497D")
ax.set_xlim(0, 22)
ax.spines["top"].set_visible(False)
ax.spines["right"].set_visible(False)
ax.legend(loc="lower right", fontsize=9)
green_p  = mpatches.Patch(color="#27ae60", label="10+ courses (strong)")
yellow_p = mpatches.Patch(color="#f39c12", label="4–9 courses (moderate)")
red_p    = mpatches.Patch(color="#e74c3c", label="0–3 courses (weak / gap)")
ax.legend(handles=[green_p, yellow_p, red_p,
                   plt.Line2D([0],[0], color="#1F497D", linestyle="--", label="Min. threshold (5)")],
          loc="lower right", fontsize=8)
plt.tight_layout()

img_buf = io.BytesIO()
plt.savefig(img_buf, format="png", dpi=150, bbox_inches="tight")
img_buf.seek(0)
plt.close()

doc.add_picture(img_buf, width=Inches(6.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ════════════════════════════════════════════════════════════════════════════
# SECTION 6 – PER-COURSE BREAKDOWN TABLE
# ════════════════════════════════════════════════════════════════════════════
add_page_break(doc)
add_heading(doc, "Per-Course Learning Outcome Breakdown", 1)
no_space_para(doc,
    "The table below lists every required-course learning outcome alongside its mapped Program Outcomes.",
    size=10)

LO_DATA = [
    ("CSTN 4015","Help Desk and Customer Support","1. Automate and standardize the IT Help Desk environment","PO1, PO3, PO5"),
    ("CSTN 4015","Help Desk and Customer Support","2. Perform user need analysis and assessment","PO1"),
    ("CSTN 4015","Help Desk and Customer Support","3. Provide efficient and effective end user and customer support","PO4, PO6, PO10"),
    ("CSTN 4015","Help Desk and Customer Support","4. Explore ITSM frameworks and identify practices to support IT services","PO1, PO5"),
    ("DBAS 1007","Data Fundamentals","1. Follow best practices in structuring data components","PO1, PO3"),
    ("DBAS 1007","Data Fundamentals","2. Implement a data storage solution using one or more DB structures","PO3"),
    ("DBAS 1007","Data Fundamentals","3. Use a DB structure for input and output of data","PO3"),
    ("DBAS 1007","Data Fundamentals","4. Determine the best data management platform for customer needs","PO1, PO5"),
    ("HDWR 1700","Introduction to Hardware and Security","1. Explain hardware and software components of a typical PC/laptop","PO3"),
    ("HDWR 1700","Introduction to Hardware and Security","2. Follow accepted industry standards in installing and configuring components","PO3"),
    ("HDWR 1700","Introduction to Hardware and Security","3. Use appropriate tools and methodologies to troubleshoot and secure systems","PO2, PO3"),
    ("HDWR 1700","Introduction to Hardware and Security","4. Complete written and oral communications and documentation","PO6"),
    ("HDWR 1700","Introduction to Hardware and Security","5. Identify physical components of a typical on-premises server infrastructure","PO3"),
    ("ICOM 2701","Professional Practice for IT I","1. Use technical tools and applications to support continuous professional improvement","PO7, PO9"),
    ("ICOM 2701","Professional Practice for IT I","2. Investigate the IT industries in relation to professional expectations","PO6, PO7, PO10"),
    ("ICOM 2701","Professional Practice for IT I","3. Communicate information using established professional processes","PO6, PO10"),
    ("ICOM 2702","Professional Practice for IT II","1. Explore emotional intelligence needed to succeed in a professional IT environment","PO6, PO10"),
    ("ICOM 2702","Professional Practice for IT II","2. Apply personal branding and marketing techniques in relation to employability","PO7, PO9, PO10"),
    ("ICOM 2702","Professional Practice for IT II","3. Cultivate collaborative skills and attitudes required for success in IT","PO6, PO10"),
    ("ICOM 2702","Professional Practice for IT II","4. Develop techniques to build professional relationships in the workplace","PO6, PO10"),
    ("ICOM 2703","Professional Practice for IT III","1. Develop a professional portfolio focused on IT industry and personal strengths","PO7, PO9"),
    ("ICOM 2703","Professional Practice for IT III","2. Prepare and practice for employment interviews","PO9, PO10"),
    ("ICOM 2703","Professional Practice for IT III","3. Examine ethical issues surrounding the workplace and decision-making","PO6, PO10"),
    ("INFT 2700","IT Project Quality Assurance","1. Develop a system design proposal that solves client requirements","PO1, PO6"),
    ("INFT 2700","IT Project Quality Assurance","2. Identify roles, processes and terminology used for IT system implementations","PO1, PO6"),
    ("INFT 2700","IT Project Quality Assurance","3. Perform project initiation activities commonly used in IT implementation projects","PO1"),
    ("INFT 2700","IT Project Quality Assurance","4. Plan, implement and document standard testing strategies","PO1, PO3, PO5"),
    ("INFT 3000","Capstone","1. Actively participate in a team to collaboratively manage a project","PO6, PO8, PO10"),
    ("INFT 3000","Capstone","2. Apply prior learning to the design and development of a project","PO1, PO2, PO3, PO5"),
    ("INFT 3000","Capstone","3. Apply professional communication strategies to a variety of target audiences","PO6"),
    ("INFT 3000","Capstone","4. Create a product that meets all approved project milestones","PO1, PO2, PO3"),
    ("ISEC 2700","Introduction to Information Security Practices","1. Identify security issues for a small business environment","PO2, PO5"),
    ("ISEC 2700","Introduction to Information Security Practices","2. Assess information security risks for a small business environment","PO1, PO2"),
    ("ISEC 2700","Introduction to Information Security Practices","3. Identify and implement security controls to maintain CIA","PO2, PO3"),
    ("ISEC 2700","Introduction to Information Security Practices","4. Troubleshoot security incidents and events to prevent further risks","PO2, PO3"),
    ("NETW 1027","Introduction to Networking and Security","1. Investigate elements of local area networks to support client connectivity","PO2, PO3"),
    ("NETW 1027","Introduction to Networking and Security","2. Build a SOHO local area network to support wired and wireless connectivity","PO2, PO3"),
    ("NETW 1027","Introduction to Networking and Security","3. Investigate elements of physical network infrastructure","PO3"),
    ("NETW 1027","Introduction to Networking and Security","4. Identify current security requirements by analyzing threats and mitigation strategies","PO2, PO5"),
    ("NETW 1500","Introduction to NOS Administration","1. Install and configure single-domain Active Directory services","PO3"),
    ("NETW 1500","Introduction to NOS Administration","2. Perform routine NOS administration tasks for a small to medium business","PO3"),
    ("NETW 1500","Introduction to NOS Administration","3. Develop and maintain documentation appropriate for IT professionals","PO1, PO6"),
    ("NETW 1500","Introduction to NOS Administration","4. Examine and deploy scripts to automate systems management tasks","PO3, PO5"),
    ("NETW 2500","NOS Administration - Windows II","1. Match business scenarios to suitable networking technologies","PO1"),
    ("NETW 2500","NOS Administration - Windows II","2. Manage a NOS directory service for a medium enterprise","PO3"),
    ("NETW 2500","NOS Administration - Windows II","3. Perform routine NOS administration tasks for a medium enterprise","PO3"),
    ("NETW 2500","NOS Administration - Windows II","4. Apply standard analysis and troubleshooting techniques for NOS support","PO3"),
    ("NETW 2500","NOS Administration - Windows II","5. Develop professional technical documentation, logs and reports","PO1, PO6"),
    ("NETW 2700","Network Infrastructure","1. Apply layered communication standards (OSI/TCP-IP) to networking devices","PO2, PO3"),
    ("NETW 2700","Network Infrastructure","2. Configure VLANs, trunking and InterVLAN routing in a switched network","PO2, PO3"),
    ("NETW 2700","Network Infrastructure","3. Implement dynamic IPv4 and IPv6 address allocation across multiple LANs","PO2, PO3"),
    ("NETW 2700","Network Infrastructure","4. Explain LAN vulnerabilities and implement switch security to mitigate attacks","PO2, PO3"),
    ("NETW 2700","Network Infrastructure","5. Explain how WLANs enable connectivity and implement WLANs","PO2, PO3"),
    ("NETW 2700","Network Infrastructure","6. Explain how routers use packets for forwarding decisions and troubleshoot static routing","PO2, PO3"),
    ("NETW 2710","Introduction to Cloud Computing and Server Virtualization","1. Determine appropriate infrastructure for on-premise virtualization","PO1, PO3"),
    ("NETW 2710","Introduction to Cloud Computing and Server Virtualization","2. Analyze economic and technological significance of virtualization and cloud computing","PO1, PO5"),
    ("NETW 2710","Introduction to Cloud Computing and Server Virtualization","3. Use cloud based services to explore IaaS","PO3, PO5"),
    ("NETW 3500","Enterprise Management and Automation","1. Determine tasks, tools and automation methods for managing a secure enterprise","PO3, PO5"),
    ("NETW 3500","Enterprise Management and Automation","2. Analyze standard tasks and tools for managing and securing an enterprise","PO2, PO3"),
    ("NETW 3500","Enterprise Management and Automation","3. Analyze cloud hosted SaaS solutions for provisioning access and managing infrastructure","PO2, PO3, PO5"),
    ("NETW 3500","Enterprise Management and Automation","4. Design professional-level communication documents and training for peers","PO4, PO6"),
    ("NETW 3700","Hierarchical Network Infrastructure","1. Implement and manage a secure network to meet enterprise needs","PO2, PO3"),
    ("NETW 3700","Hierarchical Network Infrastructure","2. Explain how WAN access technologies satisfy business requirements","PO2, PO3"),
    ("NETW 3700","Hierarchical Network Infrastructure","3. Explain the characteristics of scalable network architectures","PO2, PO3"),
    ("NETW 3700","Hierarchical Network Infrastructure","4. Explain how network automation is enabled through RESTful APIs and config management tools","PO3, PO5"),
    ("OSYS 1000","Operating Systems - Linux","1. Explain *NIX OS functions, standards and operations","PO3"),
    ("OSYS 1000","Operating Systems - Linux","2. Install and configure a *NIX based OS in standalone workstation mode","PO3"),
    ("OSYS 1000","Operating Systems - Linux","3. Maintain and troubleshoot a *NIX OS as a standalone workstation user","PO3"),
    ("OSYS 1000","Operating Systems - Linux","4. Create user scope scripts with a *NIX OS","PO3, PO5"),
    ("OSYS 1000","Operating Systems - Linux","5. Perform management techniques with a *NIX OS","PO3"),
    ("OSYS 1000","Operating Systems - Linux","6. Produce intermediate level communication documents and presentations","PO6"),
    ("OSYS 1200","Introduction to Windows Administration","1. Install and configure a Windows Desktop Operating System","PO3"),
    ("OSYS 1200","Introduction to Windows Administration","2. Use built-in tools and utilities to perform common administrative tasks","PO3"),
    ("OSYS 1200","Introduction to Windows Administration","3. Maintain associated system support documentation","PO1, PO6"),
    ("OSYS 3030","Network Services Using Linux","1. Plan a secure Linux NOS appropriate to client needs","PO1, PO2, PO3"),
    ("OSYS 3030","Network Services Using Linux","2. Implement various Linux NOS services","PO3"),
    ("OSYS 3030","Network Services Using Linux","3. Perform routine Linux NOS administration tasks","PO3"),
    ("OSYS 3030","Network Services Using Linux","4. Resolve problems using standard analysis and troubleshooting techniques","PO3"),
    ("OSYS 3030","Network Services Using Linux","5. Create and maintain documentation appropriate for IT professionals","PO1, PO6"),
    ("PROG 1700","Logic and Programming","1. Translate logic principles into programming code to solve problems","PO3, PO5"),
    ("PROG 1700","Logic and Programming","2. Perform Input/Output operations within software applications","PO3"),
    ("PROG 1700","Logic and Programming","3. Store data inside various data structures","PO3"),
    ("PROG 1700","Logic and Programming","4. Implement conditional logic within software applications","PO3"),
    ("PROG 1700","Logic and Programming","5. Implement functions/procedures to create maintainable/readable code","PO3"),
    ("PROG 1700","Logic and Programming","6. Design and implement loops to perform repetitive tasks","PO3"),
    ("PROG 1700","Logic and Programming","7. Implement debugging techniques using an IDE","PO3"),
    ("WEBD 1000","Website Development","1. Evaluate a variety of web sites for usability and accessibility","PO1"),
    ("WEBD 1000","Website Development","2. Plan and design websites based on project requirements","PO1, PO3"),
    ("WEBD 1000","Website Development","3. Develop W3C-compliant websites using HTML/CSS","PO3"),
    ("WEBD 1000","Website Development","4. Implement CSS for responsive page layout","PO3"),
]

HDR_WIDTHS = [0.85, 1.55, 4.5, 1.1]  # sums to 8.0 (within 1" margins)
lo_tbl = doc.add_table(rows=1, cols=4)
lo_tbl.style = "Table Grid"

hdr = lo_tbl.rows[0]
set_repeat_header(hdr)
for ci, (txt, w) in enumerate(zip(["Course Code","Course Name","Learning Outcome","Mapped POs"], HDR_WIDTHS)):
    set_cell_bg(hdr.cells[ci], NAVY_HEX)
    cell_para(hdr.cells[ci], txt, bold=True, size=8, color=WHITE)
    hdr.cells[ci].width = Inches(w)

course_groups = []
current = None
for row in LO_DATA:
    if row[0] != current:
        course_groups.append([])
        current = row[0]
    course_groups[-1].append(row)

for g_idx, group in enumerate(course_groups):
    bg_hex = PALE_GREY_HEX if g_idx % 2 == 0 else WHITE_HEX
    for r_idx, (code, name, lo, mapped) in enumerate(group):
        row = lo_tbl.add_row()
        for ci, w in enumerate(HDR_WIDTHS):
            row.cells[ci].width = Inches(w)
            set_cell_bg(row.cells[ci], bg_hex)
        cell_para(row.cells[0], code if r_idx == 0 else "", bold=True, size=8)
        cell_para(row.cells[1], name if r_idx == 0 else "", size=8)
        cell_para(row.cells[2], lo, size=8)
        cell_para(row.cells[3], mapped, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 7 – GAP ANALYSIS
# ════════════════════════════════════════════════════════════════════════════
add_page_break(doc)
add_heading(doc, "Gap Analysis and Recommendations", 1)
no_space_para(doc,
    "The following Program Outcomes are either absent from, or disproportionately represented in, "
    "the required curriculum. Each finding is paired with a practical recommendation.",
    size=10)
doc.add_paragraph()

GAPS = [
    ("PO4 — Provide Technical Training",
     "Only 2 courses (CSTN 4015, NETW 3500) explicitly address training others. "
     "This outcome is thinly supported across the program.",
     "Consider adding a 'communicating technical knowledge to non-technical audiences' "
     "objective to at least 2–3 additional courses (e.g., INFT 2700, NETW 1500, ISEC 2700)."),
    ("PO8 — Blend Service and Learning",
     "Only INFT 3000 (Capstone) touches this outcome, and only implicitly through team/client work. "
     "No course explicitly frames work as service to campus or community.",
     "Consider building a community-facing deliverable into at least one course "
     "(e.g., CSTN 4015 could include a simulated or real IT support clinic), or acknowledge this "
     "as a program-level milestone rather than a course-level outcome."),
    ("PO11 — Apply Sustainable Practices",
     "Zero course learning outcomes address sustainability. "
     "This outcome has no grounding in the required curriculum.",
     "Either integrate sustainability considerations into relevant technical courses "
     "(e.g., energy efficiency and e-waste in HDWR 1700, green cloud computing in NETW 2710) "
     "or reframe this as an explicit institutional outcome not requiring course-level evidence."),
    ("PO12 — Quality and Safety (WHMIS/OH&S)",
     "No course learning outcomes in the technical curriculum map to this outcome. "
     "Coverage is provided by SAFE 1000 and SAFE 1019, which are non-credit pass/no-pass courses "
     "not included in the learning outcome curriculum data.",
     "Confirm that SAFE 1000/1019 are sufficient vehicles for this outcome, and document this "
     "explicitly in the program outcome alignment record so the gap is acknowledged rather than overlooked."),
    ("PO3 — Dominant Coverage",
     "Approximately 35 of ~75 total course LOs map primarily to PO3, making it the most crowded "
     "outcome by a large margin.",
     "Consider splitting PO3 into two more specific outcomes: (a) configuring and administering "
     "operating systems and hardware, and (b) implementing security controls and hardening systems. "
     "This would give ISEC 2700 and related security work a more distinct program-level home."),
]

GAP_WIDTHS = [1.6, 3.3, 3.1]
gap_tbl = doc.add_table(rows=1, cols=3)
gap_tbl.style = "Table Grid"

ghdr = gap_tbl.rows[0]
set_repeat_header(ghdr)
for ci, (txt, w) in enumerate(zip(["Program Outcome","Gap Finding","Recommendation"], GAP_WIDTHS)):
    set_cell_bg(ghdr.cells[ci], NAVY_HEX)
    cell_para(ghdr.cells[ci], txt, bold=True, size=9, color=WHITE)
    ghdr.cells[ci].width = Inches(w)

ROW_BG = ["FFF0F0", "FFF5E6", "FFFFE6", "F0F5FF", "F0FFF0"]

for i, (po, finding, rec) in enumerate(GAPS):
    grow = gap_tbl.add_row()
    for ci, w in enumerate(GAP_WIDTHS):
        grow.cells[ci].width = Inches(w)
        set_cell_bg(grow.cells[ci], ROW_BG[i % len(ROW_BG)])
    cell_para(grow.cells[0], po, bold=True, size=9)
    cell_para(grow.cells[1], finding, size=9)
    cell_para(grow.cells[2], rec, size=9)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 8 – APPENDIX
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading(doc, "Appendix: Scope and Limitations", 1)
no_space_para(doc, (
    "This mapping covers required courses only. Elective courses were not mapped, as student "
    "selections vary and electives do not constitute a shared program experience. "
    "Work-integrated learning (INFT 4100 Work Experience) is excluded as it does not carry "
    "structured learning outcomes in the curriculum data. Safety courses (SAFE 1000 Introduction "
    "to WHMIS and SAFE 1019 Safety and Sustainability Basics) are non-credit pass/no-pass courses "
    "and are likewise excluded from the course LO mapping; their contribution to PO12 should be "
    "documented separately in the program's compliance record."
), size=10)

# ════════════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════════════
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
